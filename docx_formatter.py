import io
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.section import _Header
from docx.shared import Cm, Pt


def normalize_docx_layout(input_docx_path: str, output_docx_path: str, config: dict) -> None:
    """
    Normalize a messy .docx layout and write a cleaned version to output_docx_path.

    Args:
        input_docx_path: Path to the source .docx file.
        output_docx_path: Path to save the normalized .docx file.
        config: Layout configuration dictionary, for example:
            {
                "margin_top_cm": 2.0,
                "margin_bottom_cm": 2.0,
                "margin_left_cm": 2.5,
                "margin_right_cm": 2.5,
                "font_name": "Times New Roman",
                "font_size_pt": 12,
                "line_spacing": 1.15,
                "alignment": "justify",
                "header_text": "Nama Dokumen - Normalisasi Otomatis",
                "header_font_name": "Arial",
                "header_font_size_pt": 9,
            }

    This function preserves bold/italic run-level formatting and keeps Heading
    styles using their default/relative sizes.
    """

    document = Document(input_docx_path)

    # 1. Terapkan margin global pada setiap section.
    _apply_margins_to_all_sections(document, config)

    # 2. Terapkan font, spasi, dan alignment pada semua paragraf dan tabel.
    _normalize_document_text(document, config)

    # 3. Hapus paragraf kosong yang hanya berisi garis baru atau whitespace.
    _clean_empty_paragraphs(document)

    # 4. Tambahkan header kustom di section pertama.
    _insert_custom_header(document, config)

    # 5. Inject nomor halaman dinamis menggunakan OXML di footer.
    _inject_dynamic_page_numbering(document)

    Path(output_docx_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx_path)


def _apply_margins_to_all_sections(document: Document, config: dict) -> None:
    """Set margins for every section in the document."""
    margin_map = {
        "top": config.get("margin_top_cm", 2.0),
        "bottom": config.get("margin_bottom_cm", 2.0),
        "left": config.get("margin_left_cm", 2.5),
        "right": config.get("margin_right_cm", 2.5),
    }

    for section in document.sections:
        section.top_margin = Cm(margin_map["top"])
        section.bottom_margin = Cm(margin_map["bottom"])
        section.left_margin = Cm(margin_map["left"])
        section.right_margin = Cm(margin_map["right"])


def _normalize_document_text(document: Document, config: dict) -> None:
    """Apply font settings, line spacing, and paragraph alignment globally.

    PENGECUALIAN PENTING: paragraf di halaman COVER dan halaman IDENTITAS
    BUKU (semua yang ada SEBELUM heading "KATA PENGANTAR") sengaja TIDAK
    ditimpa alignment-nya. Halaman-halaman itu didesain center-aligned di
    template, dan kalau alignment isi naskah (biasanya "justify") diterapkan
    global ke SEMUA paragraf, halaman cover/identitas ikut berubah jadi rata
    kiri-kanan padahal seharusnya tetap center.
    """
    font_name = config.get("font_name")
    font_size_pt = config.get("font_size_pt")
    line_spacing = config.get("line_spacing")
    alignment = config.get("alignment")

    alignment_map = {
        "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
        "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
        "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
        "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
    }

    in_front_matter = True

    for paragraph in _iter_all_paragraphs(document):
        if _is_field_paragraph(paragraph):
            continue

        if in_front_matter and paragraph.text.strip().upper() == "KATA PENGANTAR":
            in_front_matter = False

        is_heading = _is_heading_style(paragraph)

        if line_spacing is not None and not in_front_matter:
            paragraph.paragraph_format.line_spacing = line_spacing

        if alignment and not in_front_matter:
            paragraph.alignment = alignment_map.get(alignment.lower(), paragraph.alignment)

        # Samakan jarak antar-paragraf isi naskah (BUKAN heading, supaya
        # space_before/after 18pt/12pt yang sengaja dipasang untuk judul bab
        # di mailmerge.py tidak ketimpa jadi 0). Tanpa ini, paragraf yang
        # datang dari naskah sumber dengan gaya beda-beda bisa punya jarak
        # antar-paragraf tidak konsisten walau line_spacing-nya sudah sama.
        if not in_front_matter and not is_heading:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

        for run in paragraph.runs:
            if font_name:
                run.font.name = font_name
            if font_size_pt is not None and not is_heading:
                run.font.size = Pt(font_size_pt)

        # Pastikan tabel juga mempertahankan styling di dalam sel.
        if paragraph._p.getparent().tag == qn("w:tc"):
            continue


def _is_heading_style(paragraph) -> bool:
    """Return True when the paragraph uses a Heading style."""
    style_name = ""
    if paragraph.style is not None:
        style_name = paragraph.style.name or ""
    return style_name.startswith("Heading")


def _iter_all_paragraphs(document: Document):
    """Yield all paragraphs from the document body and all tables."""
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table):
    """Recursively yield paragraphs inside a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                yield paragraph
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _is_field_paragraph(paragraph) -> bool:
    """Return True when the paragraph contains a Word field like TOC."""
    try:
        xml = paragraph._p.xml
    except Exception:
        return False
    return "w:instrText" in xml or "w:fldChar" in xml


def _is_section_break_paragraph(paragraph) -> bool:
    """Return True kalau paragraf ini membawa section break (w:pPr/w:sectPr).

    Paragraf semacam ini SELALU kosong teksnya (dibuat lewat
    `document.add_paragraph()` polos di mailmerge.py, lalu w:sectPr
    ditempel ke w:pPr-nya sebagai penanda "section berikutnya mulai di
    halaman ganjil baru" -- lihat `_make_odd_page_sectpr` di mailmerge.py).
    Karena teksnya kosong, tanpa pengecekan ini paragraf tsb kehitung baris
    kosong biasa oleh `_clean_empty_paragraphs` dan BISA IKUT TERHAPUS kalau
    kebetulan bersebelahan dengan baris kosong lain -- begitu paragraf
    penandanya hilang, section break-nya ikut hilang, dua section (mis.
    halaman depan & Bab 1, atau Bab N & Bab N+1) melebur balik jadi satu
    section. Akibatnya header/nomor halaman yang seharusnya beda per section
    (header dikosongkan di halaman depan, nomor halaman romawi vs arab,
    restart penomoran per bab) jadi salah/tidak konsisten, tergantung berapa
    section yang keburu melebur.
    """
    try:
        p_pr = paragraph._p.find(qn("w:pPr"))
    except Exception:
        return False
    if p_pr is None:
        return False
    return p_pr.find(qn("w:sectPr")) is not None


def _contains_image(paragraph) -> bool:
    """Return True if the paragraph has an inline/floating image (drawing) or a VML picture."""
    p_xml = paragraph._p
    drawing_tags = (
        qn("w:drawing"),
        "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline",
        "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor",
        "{urn:schemas-microsoft-com:vml}shape",
        "{urn:schemas-microsoft-com:vml}imagedata",
    )
    for tag in drawing_tags:
        if p_xml.find(f".//{tag}") is not None:
            return True
    return False


def _clean_empty_paragraphs(document: Document) -> None:
    """Rapikan paragraf kosong beruntun jadi maksimal SATU baris kosong saja.

    Paragraf berisi gambar/logo atau field Word (TOC, page number) TIDAK
    pernah dihapus. Baris kosong TUNGGAL antar bagian dibiarkan (spasi
    visual yang wajar); yang dirapikan hanya runtutan 2+ baris kosong
    berturut-turut ("spasinya banyak") supaya tidak ada jarak berlebihan.
    """
    paragraphs = list(_iter_all_paragraphs(document))
    previous_was_blank = False
    for paragraph in paragraphs:
        if _is_field_paragraph(paragraph):
            previous_was_blank = False
            continue
        if _contains_image(paragraph):
            previous_was_blank = False
            continue
        if _is_section_break_paragraph(paragraph):
            previous_was_blank = False
            continue

        is_blank = not paragraph.text or not paragraph.text.strip()
        if is_blank and previous_was_blank:
            p_element = paragraph._element
            parent = p_element.getparent()
            if parent is not None:
                parent.remove(p_element)
            continue

        previous_was_blank = is_blank


def _force_independent_definition(hdr_or_ftr) -> None:
    """Paksa header/footer ini punya definisi (part) yang BENAR-BENAR sendiri,
    bukan sekadar 'tidak linked' menurut python-docx.

    Ditemukan lewat pengetesan pakai template asli: section satu-satunya di
    template SUDAH punya elemen w:headerReference eksplisit sejak awal.
    Artinya `is_linked_to_previous` sudah False SEBELUM kita apa-apakan.
    Setiap section bab baru (hasil deep-copy sectPr template itu di
    mailmerge.py) ikut mewarisi headerReference yang SAMA (rId sama persis)
    -> python-docx menganggap semuanya "tidak linked" (masing-masing punya
    elemen reference sendiri), padahal secara fisik mereka semua menunjuk ke
    PART/isi yang SAMA PERSIS di dalam file. Assignment
    `is_linked_to_previous = False` jadi NO-OP (python-docx cuma bikin
    definisi baru kalau state-nya BERUBAH), sehingga menulis teks ke satu
    section ikut mengubah teks section lain yang berbagi part yang sama --
    termasuk halaman depan, yang jadi ikut kebawa teks header bab.

    PENTING -- kenapa TIDAK pakai trik toggle `is_linked_to_previous = True`
    lalu `= False`: assignment `True` memanggil `_drop_definition()`, yang
    membuang PART lama LEWAT `document_part.drop_header_part(rId)` --
    tindakan ini MEMBEBASKAN nomor rId itu supaya python-docx boleh
    memakainya lagi utnuk part berikutnya yang dibuat. Karena section-section
    bab di dokumen ini mewarisi rId yang SAMA PERSIS satu sama lain (hasil
    deep-copy sectPr template di mailmerge.py), begitu section pertama yang
    diproses "melepas" rId bersama itu, section BERIKUTNYA yang belum
    sempat diproses (referensinya di XML masih memuat rId lama yang sama)
    bisa "ketiban" tersambung ke part BARU yang baru saja dibuat kalau
    kebetulan part baru itu dapat alokasi nomor rId yang sama persis (rId
    yang baru saja dibebaskan). Akibatnya section pertama & section
    berikutnya berakhir menunjuk ke part yang SAMA lagi -- persis masalah
    yang ingin dihindari fungsi ini -- dan teks header/footer salah satu
    section bisa "bocor" ke section lain (mis. header judul bab ikut
    muncul di halaman depan yang seharusnya kosong).

    Solusi yang aman: lepas HANYA elemen referensinya (`w:headerReference`
    / `w:footerReference`) dari sectPr section ini -- TANPA membuang part
    lamanya dari relationship dokumen, supaya rId lama itu tidak pernah
    dibebaskan/dipakai ulang. Part lama boleh jadi masih dipakai/dirujuk
    oleh section lain yang belum diproses; itu aman dibiarkan sampai
    section itu diproses sendiri dan dapat definisi barunya sendiri. Baru
    setelah itu panggil `_add_definition()` langsung untuk membuat part
    baru + rId baru yang dijamin belum pernah dipakai (bukan hasil daur
    ulang), lalu pasang referensi barunya.
    """
    sect_pr = hdr_or_ftr._sectPr
    is_header = isinstance(hdr_or_ftr, _Header)
    hdr_ftr_type = hdr_or_ftr._hdrftr_index

    existing_reference = (
        sect_pr.get_headerReference(hdr_ftr_type)
        if is_header
        else sect_pr.get_footerReference(hdr_ftr_type)
    )
    if existing_reference is not None:
        sect_pr.remove(existing_reference)

    hdr_or_ftr._add_definition()


def _insert_custom_header(document: Document, config: dict) -> None:
    """Pasang header custom HANYA di halaman isi bab (bukan halaman depan).

    Section pertama (index 0) = halaman-halaman SEBELUM Bab 1: cover,
    identitas buku, kata pengantar, daftar isi, sinopsis. Halaman-halaman
    ini SENGAJA dibuat TANPA header judul.

    Sebelumnya header custom cuma di-set eksplisit di section[0], dan
    section-section lain (isi bab) tidak pernah diputus link header-nya
    (`is_linked_to_previous` tetap True/default) -- akibatnya, karena
    python-docx (dan Word) membuat section tanpa definisi header sendiri
    otomatis MEWARISI header dari section sebelumnya, SEMUA section isi bab
    ikut menampilkan header judul yang sama seperti front matter. Sekarang
    dibalik: front matter dikosongkan (definisi sendiri, tapi teks kosong),
    dan tiap section isi bab (index 1 dst.) diberi definisi header sendiri
    (unlinked) berisi teks judul.
    """
    header_text = config.get("header_text")
    header_font_name = config.get("header_font_name")
    header_font_size_pt = config.get("header_font_size_pt")

    sections = document.sections
    if not sections:
        return

    def _blank(hdr) -> None:
        _force_independent_definition(hdr)
        para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        para.clear()
        for extra in hdr.paragraphs[1:]:
            extra_el = extra._element
            parent = extra_el.getparent()
            if parent is not None:
                parent.remove(extra_el)

    def _fill(hdr) -> None:
        _force_independent_definition(hdr)
        para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        para.text = header_text
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.runs[0] if para.runs else para.add_run(header_text)
        if header_font_name:
            run.font.name = header_font_name
        if header_font_size_pt is not None:
            run.font.size = Pt(header_font_size_pt)

    # Section pertama: front matter, SEMUA varian header (biasa, halaman
    # genap, halaman pertama -- template ini pakai "different odd & even
    # pages") dikosongkan total.
    for variant in (sections[0].header, sections[0].even_page_header, sections[0].first_page_header):
        _blank(variant)

    if not header_text:
        return

    # Section isi bab (index 1 dst.): SEMUA varian header diisi teks judul
    # yang sama, supaya konsisten di halaman ganjil maupun genap.
    for section in sections[1:]:
        for variant in (section.header, section.even_page_header, section.first_page_header):
            _fill(variant)


def _inject_dynamic_page_numbering(document: Document) -> None:
    """Add automatic page numbering footers and reset numbering per section.

    Dokumen final bisa punya BANYAK section (bukan cuma 1-2): tiap bab baru
    disisipkan sebagai section break tersendiri (lihat mailmerge.py) supaya
    tiap bab mulai di halaman ganjil baru. Sebelumnya fungsi ini cuma
    menangani section pertama & kedua secara eksplisit — section ketiga dan
    seterusnya (bab ke-2 dst.) tidak pernah diberi nomor halaman ("belum
    berjalan"). Sekarang SEMUA section diproses:
      - Section pertama = halaman depan (cover, identitas, kata pengantar,
        daftar isi) -> angka romawi kecil (i, ii, iii, ...).
      - Section kedua dan seterusnya (isi naskah per-bab + tentang penulis)
        -> angka arab, dimulai dari 1 HANYA di section kedua; section-section
        setelahnya TIDAK di-restart supaya penomoran berlanjut wajar dan
        tidak "loncat balik ke 1" / dobel di tiap pergantian bab.
    """
    sections = document.sections

    if len(sections) == 0:
        return

    if len(sections) == 1:
        _set_section_page_number_type(sections[0], fmt="decimal", start=1)
        _set_all_footer_variants_page_number(sections[0], include_first_page=False)
        return

    # Section pertama: halaman depan, angka romawi kecil. Halaman cover
    # (first_page_footer) SENGAJA dilewati -- lihat docstring
    # `_set_all_footer_variants_page_number` -- supaya logo penerbit di
    # cover tidak ikut ketimpa field nomor halaman.
    _set_section_page_number_type(sections[0], fmt="lowerRoman", start=1)
    _set_all_footer_variants_page_number(sections[0], include_first_page=False)

    # Section kedua dst.: angka arab. Restart ke 1 cuma di section kedua
    # (awal isi naskah); section berikutnya melanjutkan penomoran otomatis
    # (start=None -> atribut w:start tidak dipasang sama sekali).
    for idx in range(1, len(sections)):
        start = 1 if idx == 1 else None
        _set_section_page_number_type(sections[idx], fmt="decimal", start=start)
        _set_all_footer_variants_page_number(sections[idx])


def _count_section_nodes(document: Document) -> int:
    """Count w:sectPr nodes in the document body XML for diagnostic purposes."""
    return len(document.element.body.findall(qn("w:sectPr")))



# Urutan resmi elemen anak <w:sectPr> menurut skema CT_SectPr (ECMA-376).
# w:pgNumType WAJIB muncul SEBELUM elemen-elemen berikut ini di dalam sectPr.
# Menaruhnya setelah salah satu dari elemen ini (mis. dengan sect_pr.append()
# begitu saja) melanggar urutan skema -> Word menolak membuka file / minta
# "repair", walaupun parser yang lebih longgar seperti LibreOffice tetap mau
# membukanya tanpa keluhan.
_SECTPR_ELEMENTS_AFTER_PGNUMTYPE = (
    "w:cols",
    "w:formProt",
    "w:vAlign",
    "w:noEndnote",
    "w:titlePg",
    "w:textDirection",
    "w:bidi",
    "w:rtlGutter",
    "w:docGrid",
    "w:printerSettings",
    "w:sectPrChange",
)


def _set_section_page_number_type(section, fmt: str, start: int | None) -> None:
    """Set the page number format for a section.

    Kalau `start` diisi (int), section ini akan RESTART penomoran mulai dari
    angka tersebut. Kalau `start` None, atribut w:start sengaja TIDAK
    dipasang sama sekali sehingga Word melanjutkan penomoran dari section
    sebelumnya (dipakai untuk bab ke-2 dst. supaya tidak restart/dobel).
    """
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgNumType"))
    if existing is not None:
        sect_pr.remove(existing)

    pg_num_type = OxmlElement("w:pgNumType")
    pg_num_type.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num_type.set(qn("w:start"), str(start))

    # Cari elemen anak pertama yang seharusnya berada SETELAH pgNumType
    # (mis. w:cols, w:titlePg, w:docGrid), lalu sisipkan pgNumType tepat
    # sebelum elemen tersebut supaya urutan skema tetap valid.
    insert_before = None
    for child in sect_pr:
        if qn_localname(child.tag) in _SECTPR_ELEMENTS_AFTER_PGNUMTYPE:
            insert_before = child
            break

    if insert_before is not None:
        insert_before.addprevious(pg_num_type)
    else:
        sect_pr.append(pg_num_type)

    try:
        _force_independent_definition(section.footer)
    except Exception:
        pass


def qn_localname(tag: str) -> str:
    """Ubah Clark-notation tag ('{namespace}local') balik jadi 'w:local' untuk dibandingkan."""
    if "}" not in tag:
        return tag
    local = tag.split("}", 1)[1]
    return f"w:{local}"


def _set_all_footer_variants_page_number(section, include_first_page: bool = True) -> None:
    """Tulis field PAGE dinamis ke varian footer section ini.

    Template buku biasanya mengaktifkan "different odd & even pages" (margin
    cermin, posisi nomor halaman gantian kiri/kanan). Kalau begitu, halaman
    genap dirender pakai `even_page_footer`, BUKAN `footer` biasa. Sebelumnya
    kode ini cuma menulis ke `section.footer`, jadi `even_page_footer` (dan
    `first_page_footer` kalau ada halaman-judul-per-bab) tetap berisi konten
    statis bawaan template -> nomor halaman di halaman genap tidak pernah
    berubah / sama terus / dobel dengan halaman lain. Sekarang ketiga varian
    ditulis field PAGE yang sama supaya penomoran selalu ikut halaman
    fisiknya, bukan konten statis lama.

    `include_first_page=False` dipakai KHUSUS untuk halaman sampul/cover
    (section pertama dokumen): halaman cover secara konvensi penerbitan
    memang tidak diberi nomor halaman, dan `first_page_footer`-nya berisi
    elemen non-teks (mis. logo penerbit) yang HARUS tetap utuh -- kalau ikut
    ditimpa field PAGE seperti varian lain, logo itu akan terhapus.
    """
    _force_independent_definition(section.footer)
    _set_footer_page_number(section.footer)

    # even_page_footer & first_page_footer HANYA relevan kalau memang dipakai
    # (Word butuh flag ini di document.settings). Aksesnya lewat python-docx
    # aman dipanggil kapan pun; kalau tidak dipakai Word akan mengabaikannya,
    # jadi lebih aman selalu menyamakan isinya drpd meninggalkan versi lama.
    variants = [getattr(section, "even_page_footer", None)]
    if include_first_page:
        variants.append(getattr(section, "first_page_footer", None))

    for variant in variants:
        if variant is None:
            continue
        try:
            _force_independent_definition(variant)
        except Exception:
            pass
        _set_footer_page_number(variant)




def _set_footer_page_number(footer) -> None:
    """Insert a centered PAGE field into the footer using OXML.

    SEMUA isi lama footer ini (paragraf lain, ATAU tabel -- banyak template
    buku memakai tabel 3 kolom kiri/tengah/kanan di footer demi perataan,
    dengan nomor halaman statis duduk di salah satu selnya) dihapus total,
    disisakan SATU paragraf kosong untuk ditulisi field PAGE. Pembersihan
    sebelumnya cuma menyisir `footer.paragraphs`, yang TIDAK menjangkau
    paragraf di dalam tabel -- makanya nomor halaman statis di dalam tabel
    footer selamat dan tetap tampil dobel bareng field yang baru.
    """
    footer_root = footer._element
    footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    kept_element = footer_paragraph._p

    for child in list(footer_root):
        if child is not kept_element:
            footer_root.remove(child)

    footer_paragraph.clear()
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = footer_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_separate)
    run._r.append(fld_end)


if __name__ == "__main__":
    sample_config = {
        "margin_top_cm": 2.0,
        "margin_bottom_cm": 2.0,
        "margin_left_cm": 3.0,
        "margin_right_cm": 3.0,
        "font_name": "Times New Roman",
        "font_size_pt": 12,
        "line_spacing": 1.15,
        "alignment": "justify",
        "header_text": "NAMA DOKUMEN - DIFORMAT OTOMATIS",
        "header_font_name": "Arial",
        "header_font_size_pt": 9,
    }

    normalize_docx_layout(
        input_docx_path="input_messy.docx",
        output_docx_path="output_normalized.docx",
        config=sample_config,
    )