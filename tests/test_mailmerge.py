import io
import os
import tempfile
import unittest

from docx import Document
from docx.oxml.ns import qn

from mailmerge import generate_book_docx


def _build_source_manuscript(chapter_titles_and_bodies):
    """Bangun naskah sumber sintetis (.docx) dengan heading 'Bab N ...' agar bisa
    dipakai sebagai fixture tanpa bergantung pada file naskah asli mana pun."""
    doc = Document()
    doc.add_heading("KATA PENGANTAR", level=1)
    doc.add_paragraph("Ini adalah kata pengantar contoh.")

    for title, body_lines in chapter_titles_and_bodies:
        doc.add_heading(title, level=1)
        for line in body_lines:
            doc.add_paragraph(line)

    doc.add_heading("TENTANG PENULIS", level=1)
    doc.add_paragraph("Biografi penulis contoh.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class MailMergeTests(unittest.TestCase):
    def test_chapter_placeholder_becomes_heading_for_toc(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "book.docx")
            generate_book_docx(
                template_path="templates/template_buku.docx",
                output_path=output_path,
                nama_penulis="Penulis Uji",
                judul_naskah="Judul Uji",
                isbn="123",
                tahun_cetak="2026",
                chapter_title="Bab 1: Pendahuluan",
                naskah_text="Isi uji naskah.",
            )

            document = Document(output_path)
            chapter_found = False

            for paragraph in document.paragraphs:
                if "Pendahuluan" in paragraph.text:
                    chapter_found = True
                    self.assertTrue(
                        paragraph.style.name.startswith("Heading")
                        or paragraph.style.name in {"1", "2", "3"},
                        msg=f"Expected heading style but got {paragraph.style.name!r}",
                    )
                    break

            self.assertTrue(chapter_found, "Chapter title was not inserted into the generated document")

    def test_multi_chapter_manuscript_writes_every_chapter_in_order(self) -> None:
        chapters = [
            ("Bab 1: Awal Cerita", ["Paragraf pertama bab satu.", "Paragraf kedua bab satu."]),
            ("Bab 2: Pertengahan", ["Paragraf pertama bab dua."]),
            ("Bab 3: Penutup", ["Paragraf pertama bab tiga.", "Paragraf kedua bab tiga."]),
        ]
        source_bytes = _build_source_manuscript(chapters)

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "book_multi.docx")
            generate_book_docx(
                template_path="templates/template_buku.docx",
                output_path=output_path,
                nama_penulis="Penulis Uji",
                judul_naskah="Judul Uji Multi Bab",
                isbn="123",
                tahun_cetak="2026",
                chapter_title="Bab 1: Awal Cerita",
                naskah_text="tidak dipakai karena source_docx_bytes tersedia",
                source_docx_bytes=source_bytes,
            )

            document = Document(output_path)

            # Semua judul bab harus ada, berurutan, masing-masing mulai halaman baru.
            found_titles = [p.text.strip() for p in document.paragraphs if p.text.strip().lower().startswith("bab")]
            expected_titles = [title for title, _ in chapters]
            self.assertEqual(found_titles, expected_titles)

            for p in document.paragraphs:
                if p.text.strip() in expected_titles:
                    self.assertTrue(
                        p.paragraph_format.page_break_before,
                        msg=f"Bab {p.text!r} seharusnya mulai di halaman baru",
                    )

            # Isi tiap bab harus ikut masuk, bukan cuma judulnya.
            full_text = "\n".join(p.text for p in document.paragraphs)
            for _, body_lines in chapters:
                for line in body_lines:
                    self.assertIn(line, full_text)

            # Field TOC otomatis harus ada, bukan lagi teks statis.
            body_xml = document.element.body.xml
            self.assertIn("TOC", body_xml)
            self.assertIn("fldChar", body_xml)

            # Word harus diarahkan untuk update field otomatis saat dibuka.
            settings_xml = document.settings.element.xml
            self.assertIn("w:updateFields", settings_xml)

    def test_manuscript_without_chapter_pattern_falls_back_to_single_chapter(self) -> None:
        """Naskah tanpa heading 'Bab N' harus tetap diperlakukan sebagai satu bab (perilaku lama)."""
        doc = Document()
        doc.add_paragraph("Paragraf naskah polos tanpa heading bab sama sekali.")
        buf = io.BytesIO()
        doc.save(buf)
        source_bytes = buf.getvalue()

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, "book_fallback.docx")
            generate_book_docx(
                template_path="templates/template_buku.docx",
                output_path=output_path,
                nama_penulis="Penulis Uji",
                judul_naskah="Judul Uji Fallback",
                isbn="123",
                tahun_cetak="2026",
                chapter_title="Bab Tunggal",
                naskah_text="teks fallback",
                source_docx_bytes=source_bytes,
            )

            document = Document(output_path)
            titles = [p.text.strip() for p in document.paragraphs if p.text.strip().lower().startswith("bab")]
            self.assertEqual(titles, ["Bab Tunggal"])


if __name__ == "__main__":
    unittest.main()