"""
Mail merge dinamis & auto-formatting untuk template buku Guepedia (.docx).
"""
import copy
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph

from docx_formatter import normalize_docx_layout
from naskah_parser import (
    extract_docx_structured,
    is_heading_line,
    is_image_marker,
    split_into_sections,
    strip_front_matter,
)

# Deteksi struktur berbasis AI (Groq) -- SEKARANG JALUR UTAMA (lihat
# docstring di ai_naskah_parser.py). Diimpor secara "lunak": kalau package
# `groq` belum terinstall atau modulnya belum ada, mailmerge.py TETAP JALAN
# NORMAL pakai parser regex/style biasa sebagai fallback (lihat
# `generate_book_docx`, parameter `use_ai_fallback`).
try:
    from ai_naskah_parser import build_sections_from_structure, detect_structure_with_ai

    _AI_PARSER_AVAILABLE = True
except ImportError:
    _AI_PARSER_AVAILABLE = False

CHAPTER_TITLE_PLACEHOLDER = "Judul Bab"
CHAPTER_BODY_PLACEHOLDER = "(Isi naskah)"

# Lebar maksimum gambar yang disisipkan dari naskah mentah ke dokumen final,
# supaya tidak melebihi margin halaman.
IMAGE_MAX_WIDTH_CM = 12.5


class MailMergeError(Exception):
    """Dilempar jika template tidak punya placeholder yang diharapkan."""


def generate_book_docx(
    template_path: str,
    output_path: str,
    nama_penulis: str,
    judul_naskah: str,
    isbn: str,
    tahun_cetak: str,
    kata_pengantar_text: str = "",
    naskah_text: str = "",
    chapter_title: str = "",
    format_config: dict | None = None,
    source_docx_bytes: bytes | None = None,
    qrcbn: str = "",
    sinopsis_text: str = "",
    tentang_penulis_text: str = "",
    use_ai_fallback: bool = True,
    groq_api_key: str | None = None,
) -> None:
    """Isi template buku dengan data dinamis, lalu terapkan format dokumen dinamis ke output akhir."""
    document = Document(template_path)

    # 1. Siapkan baris & gambar naskah, LALU tentukan strukturnya (bab,
    #    kata pengantar, sinopsis, tentang penulis) -- dilakukan DI AWAL
    #    (sebelum isi halaman depan) supaya kalau AI berhasil mendeteksi
    #    nama_penulis/judul_buku langsung dari naskah, nilai itu SEMPAT
    #    dipakai untuk mengisi halaman depan juga (bukan cuma dokumen
    #    isinya), tanpa perlu penulis mengisi ulang manual di form.
    #
    #    AI (Groq) SEKARANG JADI JALUR UTAMA (bukan cuma cadangan) selama
    #    tersedia (package `groq` terpasang + API key valid) -- parser
    #    regex/style (naskah_parser.py) cuma dipakai sebagai fallback kalau
    #    AI gagal/tidak tersedia, supaya proses tidak pernah gagal total
    #    hanya karena AI bermasalah. Alasan AI dijadikan utama: parser
    #    regex/style cuma bisa mengenali judul bab lewat pola teks "Bab N"
    #    ATAU style Word Heading -- banyak naskah nyata TIDAK cocok dengan
    #    sinyal manapun (mis. semua paragraf polos "Normal", judul bab cuma
    #    frasa pendek yang kebetulan sama dengan listing Daftar Isi-nya
    #    sendiri), dan AI juga otomatis mengenali sebutan lain untuk
    #    "Tentang Penulis" (mis. "Profil Penulis") tanpa perlu daftar
    #    sinonim manual, sekaligus mendeteksi nama penulis & judul buku
    #    dari naskah itu sendiri.
    images: dict[str, bytes] = {}
    raw_lines_for_ai: list[str] = []
    if source_docx_bytes is not None:
        source_doc = Document(io.BytesIO(source_docx_bytes))
        # Ambil gambar-nya saja di sini (murah, sekali jalan) -- kata
        # pengantar/sections dari fungsi ini TIDAK dipakai kalau AI berhasil
        # di bawah; baru dipakai ulang sebagai fallback kalau AI gagal.
        _, _, images = extract_docx_structured(source_doc)
        raw_lines_for_ai = [p.text.strip() for p in source_doc.paragraphs if p.text.strip()]
    elif naskah_text.strip():
        raw_lines_for_ai = [line.strip() for line in naskah_text.splitlines() if line.strip()]

    auto_kata_pengantar = ""
    sections: list[tuple[str, list[str]]] = []
    parsed_with_ai = False

    if use_ai_fallback and _AI_PARSER_AVAILABLE and raw_lines_for_ai:
        try:
            structure = detect_structure_with_ai(raw_lines_for_ai, api_key=groq_api_key)
            ai_kp, ai_sinopsis, ai_tentang_penulis, ai_sections = build_sections_from_structure(
                raw_lines_for_ai, structure
            )
            if ai_sections:
                # AI berhasil & mengembalikan struktur yang masuk akal --
                # pakai hasil AI sepenuhnya (termasuk Sinopsis/Tentang
                # Penulis-nya -- AI sudah mengklasifikasikannya langsung
                # berdasarkan MAKNA, bukan cuma cocokkan kata
                # "SINOPSIS"/"TENTANG PENULIS" secara harfiah, jadi otomatis
                # menangkap sebutan lain juga, mis. "PROFIL PENULIS").
                auto_kata_pengantar = ai_kp
                sections = ai_sections
                if not sinopsis_text:
                    sinopsis_text = ai_sinopsis
                if not tentang_penulis_text:
                    tentang_penulis_text = ai_tentang_penulis
                if not nama_penulis and structure.get("nama_penulis"):
                    nama_penulis = structure["nama_penulis"]
                if not judul_naskah and structure.get("judul_buku"):
                    judul_naskah = structure["judul_buku"]
                parsed_with_ai = True
        except Exception:
            # AI gagal (mis. tidak ada API key, quota habis, jaringan
            # bermasalah, respons AI bukan JSON valid) -- JANGAN sampai
            # proses generate dokumen ikut gagal total gara-gara ini.
            # Lanjut ke fallback parser regex/style di bawah.
            parsed_with_ai = False

    if not parsed_with_ai:
        if source_docx_bytes is not None:
            source_doc = Document(io.BytesIO(source_docx_bytes))
            auto_kata_pengantar, sections, images = extract_docx_structured(source_doc)
            if len(sections) == 1 and (chapter_title or not is_heading_line(sections[0][0])):
                only_title, only_body = sections[0]
                sections = [(chapter_title or only_title or "Bab 1", only_body)]
        elif naskah_text.strip():
            raw_lines = [line.strip() for line in naskah_text.splitlines() if line.strip()]
            auto_kata_pengantar, body_lines = strip_front_matter(raw_lines)
            sections = split_into_sections(body_lines, fallback_title=chapter_title)

    # 1b. Pisahkan section "SINOPSIS" / "TENTANG PENULIS" KALAU penulis
    #     kebetulan menulis sendiri bagian itu di dalam file naskahnya DAN
    #     kita jatuh ke jalur regex/style (jalur AI di atas SUDAH
    #     menanganinya sendiri secara semantik, jadi ini murni jaring
    #     pengaman untuk fallback). Tanpa ini, section semacam itu ikut
    #     disisipkan sebagai BAB BIASA di tengah buku lewat
    #     `_insert_manuscript_sections` -- makanya kata "SINOPSIS"/"TENTANG
    #     PENULIS" bisa muncul sebagai judul bab di tengah, dan isinya tidak
    #     pernah mengisi slot Sinopsis/Tentang-Penulis resmi di
    #     awal/akhir buku. Isinya dipakai sebagai FALLBACK kalau form
    #     sinopsis_text/tentang_penulis_text kosong; kalau form sudah diisi
    #     manual, section hasil ekstraksi ini cukup dibuang (dianggap
    #     duplikat) supaya tidak dobel.
    if not parsed_with_ai:
        auto_sinopsis, auto_tentang_penulis, sections = _extract_named_sections(
            sections, ("SINOPSIS", "TENTANG PENULIS")
        )
    else:
        auto_sinopsis, auto_tentang_penulis = "", ""

    # 2. Replacement Halaman Depan & Redaksi -- pakai nama_penulis/judul_naskah
    #    yang sudah (kalau perlu) ditimpa hasil deteksi AI di atas.
    simple_replacements = {
        "Nama Penulis": nama_penulis or "Nama Penulis",
        "Judul Naskah": judul_naskah or "Judul Naskah",
        "ISBN: ": f"ISBN: {isbn}" if isbn else "ISBN: ",
        "Februari 2023": tahun_cetak or "Mei 2025",
    }

    for paragraph in _iter_all_paragraphs(document):
        _apply_simple_replacements(paragraph, simple_replacements)

    if qrcbn:
        _insert_qrcbn(document, qrcbn)

    # 2. Siapkan baris & gambar naskah, buang front matter buatan penulis
    #    sendiri (Kata Pengantar & listing Daftar Isi) karena dokumen final
    #    sudah punya slotnya sendiri (placeholder Kata Pengantar + field TOC
    #    otomatis). Kalau sumbernya file .docx asli, pakai parser yang sadar
    #    STYLE Word (mengenali heading apapun teksnya, bukan cuma "Bab N")
    #    dan ikut mempertahankan gambar inline di naskah.
    images: dict[str, bytes] = {}
    if source_docx_bytes is not None:
        source_doc = Document(io.BytesIO(source_docx_bytes))
        auto_kata_pengantar, sections, images = extract_docx_structured(source_doc)
        if len(sections) == 1 and not is_heading_line(sections[0][0]):
            only_title, only_body = sections[0]
            sections = [(chapter_title or only_title or "Bab 1", only_body)]
    elif naskah_text.strip():
        raw_lines = [line.strip() for line in naskah_text.splitlines() if line.strip()]
        auto_kata_pengantar, body_lines = strip_front_matter(raw_lines)
        sections = split_into_sections(body_lines, fallback_title=chapter_title)
    else:
        auto_kata_pengantar, sections = "", []

    # 3. Isi Sinopsis (opsional — dari input manual/upload ATAU hasil generate AI di app.py)
    if sinopsis_text and sinopsis_text.strip():
        _insert_sinopsis(document, sinopsis_text)

    # 4. Isi Kata Pengantar — pakai input manual dari form kalau diisi,
    #    kalau kosong pakai yang otomatis terdeteksi dari naskah asli.
    final_kata_pengantar = kata_pengantar_text or auto_kata_pengantar
    if final_kata_pengantar:
        _replace_kata_pengantar(document, final_kata_pengantar)
    # force_page_break=True: jaga-jaga supaya Kata Pengantar tetap mulai di
    # halaman barunya sendiri walaupun (utk alasan apapun) Sinopsis kosong /
    # tidak diisi -- sebelumnya page break Kata Pengantar CUMA dipaksa lewat
    # _insert_sinopsis() (baris ~425), jadi kalau Sinopsis kosong, page break
    # itu tidak pernah di-set sama sekali.
    _style_front_matter_heading(document, "KATA PENGANTAR", force_page_break=True)

    # 5. Masukkan Isi Naskah Multi-Bab (Daftar Isi bawaan naskah sudah dibuang,
    #    begitu juga section SINOPSIS/TENTANG PENULIS bawaan naskah -- lihat
    #    langkah 2b di atas), tiap bab baru dimulai di halaman ganjil baru
    #    (section break oddPage), dan gambar dari naskah asli (kalau ada)
    #    disisipkan lagi di posisi yang sama.
    _insert_manuscript_sections(document, sections, format_config=format_config, images=images)

    # 6. Isi "Tentang Penulis" (opsional, diinput lewat form)
    if tentang_penulis_text and tentang_penulis_text.strip():
        _replace_tentang_penulis(document, tentang_penulis_text)

    # 7. Daftar Isi otomatis: taruh TEPAT SETELAH Kata Pengantar (di posisi
    #    heading "DAFTAR ISI" bawaan template), menggantikan daftar isi
    #    statis bawaan template yang nomor halamannya hardcode/pasti salah.
    _insert_toc_field(document)
    _ensure_update_fields(document)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    # 8. Terapkan Format Dinamis Global (Margin, Font, Header, Spacing, dll)
    if format_config:
        normalize_docx_layout(output_path, output_path, format_config)


def _iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph

    for section in document.sections:
        containers = [
            section.header,
            section.footer,
            section.first_page_header,
            section.first_page_footer,
            section.even_page_header,
            section.even_page_footer,
        ]
        for container in containers:
            if container is None:
                continue
            for paragraph in container.paragraphs:
                yield paragraph


def _apply_simple_replacements(paragraph: Paragraph, replacements: dict) -> None:
    text = paragraph.text
    if not text:
        return

    new_text = text
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)

    new_text = _sanitize_docx_text(new_text)

    if new_text != text:
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


def _insert_qrcbn(document: Document, qrcbn: str) -> None:
    """Sisipkan baris QRCBN tepat setelah baris ISBN di halaman francis."""
    safe_qrcbn = _sanitize_docx_text(qrcbn)
    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().startswith("ISBN:"):
            anchor_element = paragraph._p
            new_p = document.add_paragraph()
            anchor_element.addnext(new_p._p)
            new_p.text = f"QRCBN: {safe_qrcbn}" if safe_qrcbn else "QRCBN:"
            for run in paragraph.runs:
                # samakan format (font/size) dengan baris ISBN di atasnya
                if run.text.strip():
                    for new_run in new_p.runs:
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
                    break
            return


def _extract_named_sections(
    sections: list[tuple[str, list[str]]],
    names: tuple[str, ...],
) -> tuple[str, str, list[tuple[str, list[str]]]]:
    """Pisahkan bagian "SINOPSIS" / "TENTANG PENULIS" tulisan penulis sendiri
    dari daftar section isi naskah biasa hasil naskah_parser, DALAM DUA BENTUK:

    (1) SECTION TERSENDIRI -- naskah_parser sempat mendeteksinya sebagai
        section/heading sendiri (judul section == persis salah satu nama di
        `names`). Ini terjadi kalau baris itu memang dikenali sebagai heading
        oleh naskah_parser.

    (2) BARIS BIASA DI TENGAH BODY TEXT suatu bab -- ternyata BENTUK YANG
        LEBIH SERING TERJADI: penulis menulis kata "TENTANG PENULIS" /
        "SINOPSIS" sebagai paragraf polos di file .docx sumber (bukan pakai
        style Heading Word), sehingga naskah_parser (yang mendeteksi heading
        berdasarkan STYLE, bukan isi teks) TIDAK PERNAH menganggapnya sebagai
        section/judul sendiri -- baris itu ikut lolos sebagai body_lines
        biasa milik bab manapun yang kebetulan sedang berjalan saat baris itu
        muncul di naskah. Kasus (1) saja TIDAK CUKUP menangkap ini karena
        baris tsb tidak pernah jadi "judul section" untuk mulai dengan --
        makanya perlu di-scan juga di DALAM body_lines tiap section.

    Dalam bentuk (2), begitu baris penanda ("SINOPSIS" / "TENTANG PENULIS")
    ditemukan di tengah body_lines, SEMUA baris sesudahnya (lintas section,
    sampai ada penanda lain atau naskah habis) dianggap MILIK penanda
    tersebut -- karena di praktiknya bagian ini biasa ditulis penulis di
    paling akhir naskah tanpa heading baru lagi sesudahnya utk menutupnya.

    Hanya baris yang PERSIS cocok (case-insensitive, sesudah di-strip) yang
    dianggap penanda, supaya tidak salah menangkap kalimat isi bab yang
    kebetulan memuat kata "sinopsis"/"tentang penulis" di tengah kalimat.

    Return: (gabungan_baris_sinopsis, gabungan_baris_tentang_penulis, sisa_sections)
    """
    normalized_targets = {name.strip().upper() for name in names}
    collected: dict[str, list[str]] = {name: [] for name in normalized_targets}
    remaining: list[tuple[str, list[str]]] = []

    # Penanda BUKAN-target yang menandai "akhir" penampungan penanda aktif --
    # kalau tidak ada ini, begitu satu penanda aktif (mis. SINOPSIS) ketemu,
    # SEMUA baris sesudahnya (termasuk bagian lain milik naskah yg sah, mis.
    # DAFTAR PUSTAKA/bibliografi) ikut tersedot masuk jadi "isi sinopsis" --
    # padahal daftar pustaka itu bagian isi buku yang sah dan harus tetap
    # muncul sebagai body biasa, bukan ikut hilang ditelan bucket sinopsis.
    stop_markers = {
        "DAFTAR PUSTAKA", "DAFTAR REFERENSI", "REFERENSI", "BIBLIOGRAFI", "DAFTAR RUJUKAN",
    }

    # Penanda yang sedang "aktif" -- begitu ketemu baris penanda di tengah
    # body_lines, baris-baris SESUDAHNYA (termasuk lintas section berikutnya)
    # ditampung ke situ, sampai ketemu penanda lain (target ATAU stop_markers)
    # atau naskah habis.
    active_marker: str | None = None

    for title, body_lines in sections:
        title_key = (title or "").strip().upper()

        if title_key in normalized_targets:
            # Kasus (1): section ini sendiri memang persis SINOPSIS/TENTANG
            # PENULIS -- seluruh isinya milik penanda itu, judul section
            # tidak perlu diproses lagi sebagai body biasa.
            collected[title_key].extend(body_lines)
            active_marker = None
            continue

        new_body_lines: list[str] = []
        for line in body_lines:
            line_key = line.strip().upper()
            if line_key in normalized_targets:
                # Kasus (2): baris penanda ketemu DI TENGAH body_lines.
                active_marker = line_key
                continue
            if line_key in stop_markers:
                # Balik ke mode isi bab biasa -- baris ini & sesudahnya tetap
                # bagian sah dari naskah, JANGAN ikut tertelan bucket penanda.
                active_marker = None
                new_body_lines.append(line)
                continue
            if active_marker is not None:
                collected[active_marker].append(line)
            else:
                new_body_lines.append(line)

        remaining.append((title, new_body_lines))

    auto_sinopsis = "\n".join(collected.get("SINOPSIS", [])).strip()
    auto_tentang_penulis = "\n".join(collected.get("TENTANG PENULIS", [])).strip()
    return auto_sinopsis, auto_tentang_penulis, remaining


def _insert_sinopsis(document: Document, sinopsis_text: str) -> None:
    """Sisipkan halaman SINOPSIS baru sebelum halaman Kata Pengantar."""
    anchor_paragraph = None
    for paragraph in document.paragraphs:
        if "KATA PENGANTAR" in paragraph.text:
            anchor_paragraph = paragraph
            break

    if anchor_paragraph is None:
        anchor_paragraph = document.paragraphs[-1]

    lines = [_normalize_spacing(line) for line in _sanitize_docx_text(sinopsis_text).splitlines() if line.strip()]
    if not lines:
        return

    anchor_element = anchor_paragraph._p

    heading_p = document.add_paragraph()
    anchor_element.addprevious(heading_p._p)
    heading_p.text = "SINOPSIS"
    try:
        heading_p.style = document.styles["Heading 1"]
    except KeyError:
        pass
    for run in heading_p.runs:
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        run.font.bold = True
    heading_p.paragraph_format.page_break_before = True

    insert_after = heading_p._p
    for line in lines:
        body_p = document.add_paragraph()
        insert_after.addnext(body_p._p)
        insert_after = body_p._p
        body_p.text = line
        body_p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        body_p.paragraph_format.first_line_indent = Cm(0.75)

    # Kata Pengantar tetap mulai di halaman baru setelah Sinopsis
    anchor_paragraph.paragraph_format.page_break_before = True


def _sanitize_docx_text(text: str) -> str:
    """Bersihkan karakter kontrol yang tidak boleh ada di XML/Word sebelum ditulis ke DOCX."""
    if not text:
        return ""

    if not isinstance(text, str):
        text = str(text)

    sanitized_chars = []
    for char in text:
        code_point = ord(char)
        if code_point in (9, 10, 13):
            sanitized_chars.append(char)
        elif 0 <= code_point <= 31 or 127 <= code_point <= 159:
            sanitized_chars.append(" ")
        else:
            sanitized_chars.append(char)

    return "".join(sanitized_chars)


def _normalize_spacing(text: str) -> str:
    """Rapikan spasi ganda/berlebih jadi satu spasi, buang spasi awal-akhir.

    Juga buang sisa artefak nomor halaman Daftar Isi ('\\t1' dsb di ujung
    baris) yang kadang ikut kebawa kalau naskah sumbernya berupa Daftar Isi
    ketikan manual (bukan field TOC Word) yang lalu ke-copy-paste sebagai
    teks biasa.
    """
    text = _sanitize_docx_text(text)
    text = re.sub(r"\t+\d+\s*$", "", text)
    return re.sub(r"[ \t]{2,}", " ", text).strip()


def _style_front_matter_heading(
    document: Document, label: str, level: str = "Heading 1", force_page_break: bool = False
) -> None:
    """Paksa paragraf placeholder front-matter (mis. 'KATA PENGANTAR') jadi
    Heading Word beneran + warna hitam.

    Placeholder seperti ini biasanya cuma teks biasa/bold manual di template
    (bukan style Heading Word asli), jadi selama ini: (1) warnanya tidak ikut
    kena fix hitam yang dipasang utk judul bab, dan (2) tidak pernah muncul
    di Daftar Isi otomatis karena field TOC (\\o "1-3") cuma menyaring
    paragraf ber-style Heading 1-3, bukan berdasar teks/bold manual.

    `force_page_break`: kalau True, paragraf heading ini SELALU dipaksa mulai
    di halaman baru (page_break_before), apapun isi/panjang bagian sebelumnya.
    Wajib True untuk heading yang posisinya di dokumen bergantung pada konten
    dinamis di depannya (mis. "TENTANG PENULIS" yang selalu jatuh setelah
    bab terakhir -- jumlah & panjang bab beda-beda tiap naskah, jadi kalau
    tidak dipaksa, halaman terakhir bab terakhir yang kebetulan masih ada
    sisa ruang bikin "TENTANG PENULIS" nyambung di halaman yang sama alih-alih
    mulai halaman barunya sendiri).
    """
    for paragraph in document.paragraphs:
        if paragraph.text.strip().upper() != label.upper():
            continue
        try:
            paragraph.style = document.styles[level]
        except KeyError:
            pass
        if not paragraph.runs:
            paragraph.add_run(paragraph.text)
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            run.font.bold = True
        if force_page_break:
            paragraph.paragraph_format.page_break_before = True
        break


def _replace_kata_pengantar(document: Document, kata_pengantar_text: str) -> None:
    """Isi bagian Kata Pengantar sebelum Bab 1."""
    lines = [_normalize_spacing(line) for line in _sanitize_docx_text(kata_pengantar_text).splitlines() if line.strip()]
    if not lines:
        return

    for i, paragraph in enumerate(document.paragraphs):
        if "KATA PENGANTAR" in paragraph.text:
            if i + 1 < len(document.paragraphs):
                target = document.paragraphs[i + 1]
                target.text = lines[0]
                anchor_element = target._p
                parent = target._parent
                for line in lines[1:]:
                    new_element = copy.deepcopy(anchor_element)
                    anchor_element.addnext(new_element)
                    anchor_element = new_element
                    new_paragraph = Paragraph(new_element, parent)
                    new_paragraph.text = line
            break


def _replace_tentang_penulis(document: Document, tentang_penulis_text: str) -> None:
    """Isi bagian 'Tentang Penulis' di akhir buku (placeholder ada di template,
    tepat setelah heading 'TENTANG PENULIS')."""
    lines = [_normalize_spacing(line) for line in _sanitize_docx_text(tentang_penulis_text).splitlines() if line.strip()]
    if not lines:
        return

    for i, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().upper() == "TENTANG PENULIS":
            if i + 1 < len(document.paragraphs):
                target = document.paragraphs[i + 1]
                target.text = lines[0]
                anchor_element = target._p
                parent = target._parent
                for line in lines[1:]:
                    new_element = copy.deepcopy(anchor_element)
                    anchor_element.addnext(new_element)
                    anchor_element = new_element
                    new_paragraph = Paragraph(new_element, parent)
                    new_paragraph.text = line
            break


def _insert_manuscript_sections(
    document: Document,
    sections: list[tuple[str, list[str]]],
    format_config: dict | None = None,
    images: dict[str, bytes] | None = None,
) -> None:
    """Sisipkan sections (judul_bab, isi_baris) hasil naskah_parser ke placeholder bab, dengan styling dinamis."""
    images = images or {}
    target_p = None
    for paragraph in document.paragraphs:
        if CHAPTER_BODY_PLACEHOLDER in paragraph.text or CHAPTER_TITLE_PLACEHOLDER in paragraph.text:
            target_p = paragraph
            break

    if target_p is None:
        # Placeholder "Judul Bab"/"(Isi naskah)" tidak ada di template ini.
        # DULU: fallback-nya `document.add_paragraph()`, yang nambah paragraf
        # baru di paling UJUNG DOKUMEN -- kalau template sudah punya heading
        # "TENTANG PENULIS" di dekat akhir (lazim di template Guepedia), isi
        # naskah/bab jadi ke-insert SETELAH "Tentang Penulis", membalik
        # urutan yang seharusnya (Tentang Penulis harus di paling akhir,
        # setelah bab terakhir, bukan sebelum bab pertama).
        # SEKARANG: kalau placeholder tidak ketemu, coba dulu sisip TEPAT
        # SEBELUM heading "TENTANG PENULIS" (kalau ada) supaya urutan buku
        # tetap benar walau template lupa taruh placeholder-nya. Baru kalau
        # itu juga tidak ada, fallback paling akhir ke ujung dokumen seperti
        # sebelumnya (drpd gagal total).
        tentang_penulis_heading = None
        for paragraph in document.paragraphs:
            if paragraph.text.strip().upper() == "TENTANG PENULIS":
                tentang_penulis_heading = paragraph
                break

        if tentang_penulis_heading is not None:
            target_p = document.add_paragraph()
            tentang_penulis_heading._p.addprevious(target_p._p)
        else:
            target_p = document.add_paragraph()

    if not sections:
        p_element = target_p._element
        if p_element.getparent() is not None:
            p_element.getparent().remove(p_element)
        return

    # Ambil setting font & line spacing dari format_config dinamis
    font_name = format_config.get("font_name", "Times New Roman") if format_config else "Times New Roman"
    font_size = format_config.get("font_size_pt", 12) if format_config else 12
    line_spacing = format_config.get("line_spacing", 1.15) if format_config else 1.15
    heading_alignment_key = (format_config.get("heading_alignment", "left") if format_config else "left")
    heading_alignment = _ALIGNMENT_MAP.get(heading_alignment_key.lower(), WD_PARAGRAPH_ALIGNMENT.LEFT)

    # sectPr dasar dipakai sebagai cetakan section break tiap awal bab
    # (disalin dari section terakhir template supaya margin/header/footer konsisten)
    base_sect_pr = document.element.body.find(qn("w:sectPr"))

    anchor_element = target_p._p

    def _new_paragraph_after(anchor_el):
        """Buat paragraf baru lewat API python-docx (murah), lalu pindahkan
        tepat setelah anchor_el. Jauh lebih hemat memori/CPU dibanding
        copy.deepcopy per baris, yang mahal untuk naskah ratusan halaman."""
        p = document.add_paragraph()
        anchor_el.addnext(p._p)
        return p

    for section_title, body_lines in sections:
        if section_title:
            # is_bab = True untuk SETIAP judul yang lolos deteksi heading di
            # naskah_parser (pola "Bab N" klasik, angka romawi/eja, MAUPUN
            # heading hasil STYLE Word lewat extract_docx_structured) —
            # kecuali sub-heading bergaya "[...]" yang sengaja diperlakukan
            # sebagai Heading 2, bukan bab utama.
            is_sub_heading = section_title.startswith("[") and section_title.endswith("]")
            is_bab = not is_sub_heading

            if is_bab and base_sect_pr is not None:
                # Section break "mulai halaman ganjil baru" tepat sebelum judul bab.
                marker_p = _new_paragraph_after(anchor_element)
                anchor_element = marker_p._p
                marker_p_pr = marker_p._p.get_or_add_pPr()
                marker_p_pr.append(_make_odd_page_sectpr(base_sect_pr))

            p_new = _new_paragraph_after(anchor_element)
            anchor_element = p_new._p

            safe_title = _normalize_spacing(section_title)
            p_new.text = safe_title

            if is_bab:
                try:
                    p_new.style = document.styles["Heading 1"]
                except KeyError:
                    p_new.style = document.styles["Normal"]
                p_new.alignment = heading_alignment
                p_new.paragraph_format.space_before = Pt(18)
                p_new.paragraph_format.space_after = Pt(12)
                p_new.paragraph_format.page_break_before = True
                heading_font_size = font_size + 4
            elif is_sub_heading:
                try:
                    p_new.style = document.styles["Heading 2"]
                except KeyError:
                    p_new.style = document.styles["Normal"]
                p_new.alignment = heading_alignment
                p_new.paragraph_format.space_before = Pt(12)
                p_new.paragraph_format.space_after = Pt(4)
                heading_font_size = font_size + 2
            else:
                p_new.style = document.styles["Normal"]
                p_new.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                heading_font_size = font_size

            for run in p_new.runs:
                run.font.name = font_name
                run.font.size = Pt(heading_font_size)
                if is_sub_heading or is_bab:
                    run.font.bold = True
                    # Style "Heading 1"/"Heading 2" bawaan Word defaultnya
                    # pakai warna tema (biru accent). Dipaksa hitam supaya
                    # judul bab konsisten dengan warna teks isi naskah.
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

        for line in body_lines:
            if is_image_marker(line) and line in images:
                # Gambar dari naskah mentah (mis. foto/ilustrasi sisipan)
                # dipertahankan di posisi aslinya, bukan hilang begitu saja.
                p_img = _new_paragraph_after(anchor_element)
                anchor_element = p_img._p
                p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = p_img.add_run()
                try:
                    run.add_picture(io.BytesIO(images[line]), width=Cm(IMAGE_MAX_WIDTH_CM))
                except Exception:
                    # Gambar korup/format tak didukung -> lewati saja, jangan
                    # gagalkan seluruh proses generate dokumen.
                    pass
                continue

            clean_line = _normalize_spacing(line)
            if not clean_line:
                continue

            p_new = _new_paragraph_after(anchor_element)
            anchor_element = p_new._p

            p_new.text = clean_line

            # Paragraf Isi Naskah biasa (Format Dinamis).
            # space_before/space_after DIPAKSA 0: kalau tidak, paragraf baru
            # ikut jarak bawaan style "Normal" (biasanya ada space_after
            # beberapa pt), yang bikin jarak antar-paragraf jadi tidak rata
            # dibanding rapatnya baris dalam satu paragraf (line_spacing).
            # Book layout memisahkan paragraf lewat first_line_indent, bukan
            # jarak vertikal ekstra.
            p_new.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p_new.paragraph_format.first_line_indent = Cm(0.75)
            p_new.paragraph_format.line_spacing = line_spacing
            p_new.paragraph_format.space_before = Pt(0)
            p_new.paragraph_format.space_after = Pt(0)

            for run in p_new.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    # Hapus paragraf placeholder bawaan
    p_element = target_p._element
    if p_element.getparent() is not None:
        p_element.getparent().remove(p_element)


_ALIGNMENT_MAP = {
    "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
    "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
    "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
}

# Elemen sectPr yang menurut skema CT_SectPr WAJIB muncul SEBELUM w:type.
_SECTPR_LOCALNAMES_BEFORE_TYPE = {"headerReference", "footerReference", "footnotePr", "endnotePr"}


def _make_odd_page_sectpr(base_sect_pr) -> "OxmlElement":
    """Salin sectPr dasar & set w:type jadi 'oddPage' — dipakai sebagai
    penanda section break di w:pPr suatu paragraf, supaya bab berikutnya
    otomatis mulai di halaman ganjil baru (Word akan menyisipkan halaman
    kosong sendiri kalau perlu)."""
    new_sect_pr = copy.deepcopy(base_sect_pr)

    # Buang titlePg: kalau ikut disalin, halaman pertama SETIAP bab akan
    # pakai header/footer "halaman judul" (biasanya tanpa nomor halaman),
    # padahal kita mau nomor halaman tetap konsisten di halaman pembuka bab.
    for title_pg in new_sect_pr.findall(qn("w:titlePg")):
        new_sect_pr.remove(title_pg)

    for existing_type in new_sect_pr.findall(qn("w:type")):
        new_sect_pr.remove(existing_type)

    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "oddPage")

    insert_after = None
    for child in new_sect_pr:
        if _local_tag(child.tag) in _SECTPR_LOCALNAMES_BEFORE_TYPE:
            insert_after = child
    if insert_after is not None:
        insert_after.addnext(type_el)
    else:
        new_sect_pr.insert(0, type_el)

    return new_sect_pr


def _write_toc_field_into(paragraph) -> None:
    """Tulis field TOC otomatis ke sebuah paragraf yang sudah ada.

    PENTING: w:fldChar / w:instrText harus jadi anak <w:r> (run), BUKAN anak
    langsung <w:p> (paragraf) — kalau tidak, Word menganggap file corrupt dan
    minta "repair" saat dibuka. Ikuti pola begin -> separate -> end yang sama
    dengan _set_footer_page_number di docx_formatter.py.
    """
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instr_text)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _insert_toc_field(document: Document) -> None:
    """Sisipkan field Daftar Isi (TOC) OTOMATIS tepat SETELAH heading
    'DAFTAR ISI' bawaan template — yaitu di awal buku, setelah Kata
    Pengantar, sesuai posisi aslinya di template. Sebelumnya field ini
    ditambahkan di paragraf paling akhir dokumen (setelah Tentang Penulis),
    yang salah tempat.

    Daftar isi statis bawaan template (baris-baris berstyle 'toc N' persis
    setelah heading, isinya nomor halaman hardcode yang pasti meleset begitu
    isi naskah berubah) ikut dibuang karena digantikan field otomatis ini.
    """
    anchor = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip().upper() == "DAFTAR ISI":
            anchor = paragraph
            break

    if anchor is None:
        # Fallback: heading "DAFTAR ISI" tidak ditemukan di template --
        # taruh di akhir dokumen seperti perilaku lama, drpd tidak ada TOC sama sekali.
        paragraph = document.add_paragraph()
        _write_toc_field_into(paragraph)
        return

    # Buang baris daftar isi statis bawaan template (style 'toc N') yang ada
    # tepat setelah heading "DAFTAR ISI" — boleh diselingi baris kosong,
    # jadi baris kosong "ditunda" dulu dan baru ikut dibuang kalau ternyata
    # diikuti entri toc (supaya baris kosong yang BUKAN bagian dari listing
    # toc tidak ikut kehapus).
    next_el = anchor._p.getnext()
    pending_blanks = []
    while next_el is not None:
        following_el = next_el.getnext()
        temp_paragraph = Paragraph(next_el, anchor._parent)
        style_name = (temp_paragraph.style.name if temp_paragraph.style is not None else "") or ""
        text = temp_paragraph.text.strip()

        if style_name.lower().startswith("toc"):
            for blank_el in pending_blanks:
                blank_el.getparent().remove(blank_el)
            pending_blanks = []
            next_el.getparent().remove(next_el)
            next_el = following_el
            continue

        if not text:
            pending_blanks.append(next_el)
            next_el = following_el
            continue

        break

    toc_paragraph = document.add_paragraph()
    anchor._p.addnext(toc_paragraph._p)
    _write_toc_field_into(toc_paragraph)


def _local_tag(tag: str) -> str:
    """Ubah Clark-notation tag ('{namespace}local') jadi nama lokalnya saja."""
    return tag.split("}", 1)[1] if "}" in tag else tag


# Nama elemen (localname, tanpa prefix) yang menurut skema CT_Settings WAJIB
# muncul SETELAH w:updateFields. Dipakai untuk mencari titik sisip yang benar
# supaya urutan skema tidak dilanggar (append() mentah ke akhir <w:settings>
# bisa menaruh updateFields SETELAH elemen seperti w:compat/w:rsids/mathPr,
# yang membuat Word menganggap file corrupt).
_SETTINGS_LOCALNAMES_AFTER_UPDATEFIELDS = {
    "hdrShapeDefaults", "footnotePr", "endnotePr", "compat", "docVars",
    "rsids", "mathPr", "attachedSchema", "themeFontLang", "clrSchemeMapping",
    "doNotIncludeSubdocsInStats", "doNotAutoCompressPictures", "forceUpgrade",
    "captions", "readModeInkLockDown", "smartTagType", "schemaLibrary",
    "shapeDefaults", "doNotEmbedSmartTags", "decimalSymbol", "listSeparator",
    "doNotDemarcateInvalidXml", "doNotValidateAgainstSchema",
}


def _ensure_update_fields(document: Document) -> None:
    """Pastikan Word diperintahkan meng-update field TOC saat dokumen dibuka."""
    settings_element = document.settings.element
    existing = settings_element.find(qn("w:updateFields"))
    if existing is not None:
        return

    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "1")

    insert_before = None
    for child in settings_element:
        if _local_tag(child.tag) in _SETTINGS_LOCALNAMES_AFTER_UPDATEFIELDS:
            insert_before = child
            break

    if insert_before is not None:
        insert_before.addprevious(update_fields)
    else:
        settings_element.append(update_fields)